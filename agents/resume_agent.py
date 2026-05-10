"""
Resume Agent - Tailors resumes to specific job postings.
Uses Claude Sonnet 4.5 for high-quality rewriting.
"""
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.claude_client import ClaudeClient


class ResumeAgent:
    """Agent for generating tailored resumes based on job requirements."""

    def __init__(self, profile_path: str = "data/candidate_profile.json"):
        self.profile_path = profile_path
        self.candidate_profile = self._load_profile()
        self.client = ClaudeClient(model="claude-opus-4-5")

    def _load_profile(self) -> Dict[str, Any]:
        """Load candidate profile."""
        with open(self.profile_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def generate_tailored_resume(
        self,
        job: Dict[str, Any],
        match_report: Dict[str, Any],
        output_dir: str = "artifacts/resumes"
    ) -> Dict[str, str]:
        """Generate a tailored resume for a specific job."""
        print(f"\nGenerating tailored resume for: {job['title']} at {job['company']}")
        print("-" * 80)

        resume_content = self._generate_resume_content(job, match_report)

        job_hash = job.get('job_hash', 'unknown')
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        md_file = output_path / f"resume_{job_hash}.md"
        docx_file = output_path / f"resume_{job_hash}.docx"
        changelog_file = output_path / f"resume_{job_hash}_changelog.md"

        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(resume_content['resume_markdown'])

        with open(changelog_file, 'w', encoding='utf-8') as f:
            f.write(resume_content['change_log'])

        self._create_docx(resume_content['resume_markdown'], docx_file)

        print(f"  [OK] Markdown: {md_file}")
        print(f"  [OK] DOCX: {docx_file}")
        print(f"  [OK] Changelog: {changelog_file}")

        return {
            'markdown': str(md_file),
            'docx': str(docx_file),
            'changelog': str(changelog_file)
        }

    def _generate_resume_content(self, job: Dict[str, Any], match_report: Dict[str, Any]) -> Dict[str, str]:
        """Generate tailored resume content using Claude."""
        prompt = self._build_resume_prompt(job, match_report)
        system_prompt = self._build_resume_system_prompt()

        response = self.client.generate(
            prompt=prompt,
            system=system_prompt,
            max_tokens=4096
        )

        parts = response.split("---CHANGELOG---")

        resume_markdown = parts[0].strip()
        change_log = parts[1].strip() if len(parts) > 1 else "No changelog provided."

        return {
            'resume_markdown': resume_markdown,
            'change_log': change_log
        }

    def _build_resume_system_prompt(self) -> str:
        """Build system prompt for resume generation."""
        return """You are an expert resume writer specializing in AI/ML roles.

Your task is to tailor a resume to a specific job posting while maintaining 100% honesty.

STRICT RULES:
1. Use ONLY information from the candidate profile provided - NO fabrication
2. Rewrite experience bullets to naturally mirror job description keywords
3. Prioritize most relevant projects/experience to the top
4. Address skill gaps honestly (e.g., "exposure to X through Y project")
5. Target 550-650 words MAXIMUM - prioritize quality over quantity
6. If resume approaches 700 words, cut least relevant bullets
7. Use strong action verbs and quantifiable achievements where available
8. Maintain professional formatting

WORD COUNT TARGET: 550-650 words (absolute max 700)
- Cut weak bullets, keep strong ones
- Prioritize technical specifics over generic statements
- Trim wordy descriptions

FORMAT:
Output your response in TWO sections separated by "---CHANGELOG---":

SECTION 1: Complete resume in Markdown format (550-650 words)
SECTION 2: Detailed change log explaining what was changed and why

Example structure:
```
# Rudra Joshi
[Contact info]

## Professional Summary
[Tailored 2-3 sentence summary]

## Skills
[Categorized skills, prioritized by relevance]

## Education
[Education details]

## Projects / Experience
[Most relevant projects first, bullets rewritten to match JD keywords]

---CHANGELOG---

## Resume Tailoring Changes

1. **Professional Summary**: Emphasized [X] because the role requires [Y]
2. **Skills Section**: Moved [skill] to top as it appears in required skills
3. **Project 1**: Reworded bullet to highlight [keyword from JD]
...
```

Be strategic about what to emphasize, but never invent experience."""

    def _build_resume_prompt(self, job: Dict[str, Any], match_report: Dict[str, Any]) -> str:
        """Build prompt for resume generation."""
        profile_json = json.dumps(self.candidate_profile, indent=2)

        return f"""Tailor a resume for this job posting based on the candidate profile.

JOB DETAILS:
Title: {job['title']}
Company: {job['company']}
Location: {job['location']}
Seniority: {job.get('seniority_level', 'N/A')}
Employment Type: {job.get('employment_type', 'N/A')}

Required Skills: {', '.join(job.get('required_skills', []))}
Nice-to-Have Skills: {', '.join(job.get('nice_to_have_skills', []))}

Responsibilities:
{chr(10).join(f'- {r}' for r in job.get('responsibilities', []))}

MATCH ANALYSIS:
Score: {match_report['score']}/100
Matching Skills: {', '.join(match_report['matching_skills'])}
Missing Skills: {', '.join(match_report['missing_skills'])}

Reasons: {match_report['reasons']}

CANDIDATE PROFILE:
{profile_json}

TASK:
Create a tailored resume that:
1. Emphasizes the {len(match_report['matching_skills'])} matching skills
2. Addresses the {len(match_report['missing_skills'])} missing skills honestly where possible
3. Rewords experience bullets to naturally include keywords from required skills
4. Prioritizes most relevant projects/experience for this specific role
5. Uses ONLY real information from the profile above

CRITICAL WORD COUNT REQUIREMENT:
Final resume MUST be between 500-650 words total.

Before returning your response:
1. Count all words in your resume markdown (excluding "---CHANGELOG---" and changelog sections)
2. If word count is over 650 words:
   - Identify the weakest bullet point in the least relevant section
   - Remove that bullet point
   - Recount the words
   - Repeat until word count is between 500-650 words
3. Do NOT return a resume that exceeds 650 words or is below 500 words

Remember: Output resume markdown (500-650 words), then "---CHANGELOG---", then detailed changelog."""

    def _create_docx(self, markdown_content: str, output_file: Path):
        """Convert markdown resume to DOCX format."""
        doc = Document()

        # Set narrow margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(36)
            section.bottom_margin = Pt(36)
            section.left_margin = Pt(54)
            section.right_margin = Pt(54)

        lines = markdown_content.split('\n')

        for line in lines:
            line = line.strip()

            if not line:
                continue

            if line.startswith('# '):
                # Main heading (name)
                para = doc.add_paragraph(line[2:])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                run.font.size = Pt(20)
                run.font.bold = True

            elif line.startswith('## '):
                # Section heading
                para = doc.add_paragraph(line[3:])
                run = para.runs[0]
                run.font.size = Pt(14)
                run.font.bold = True
                para.space_before = Pt(12)
                para.space_after = Pt(6)

            elif line.startswith('### '):
                # Sub-heading
                para = doc.add_paragraph(line[4:])
                run = para.runs[0]
                run.font.size = Pt(12)
                run.font.bold = True

            elif line.startswith('- '):
                # Bullet point
                para = doc.add_paragraph(line[2:], style='List Bullet')
                run = para.runs[0]
                run.font.size = Pt(11)

            elif line.startswith('* '):
                # Bullet point (alternative)
                para = doc.add_paragraph(line[2:], style='List Bullet')
                run = para.runs[0]
                run.font.size = Pt(11)

            else:
                # Normal paragraph
                para = doc.add_paragraph(line)
                run = para.runs[0]
                run.font.size = Pt(11)

        doc.save(output_file)


def main():
    """Example usage."""
    agent = ResumeAgent()

    # Load a shortlisted job
    with open('data/jobs_shortlisted.json', 'r') as f:
        shortlisted = json.load(f)

    if shortlisted:
        top_job = shortlisted[0]
        files = agent.generate_tailored_resume(
            job=top_job['job'],
            match_report=top_job['match']
        )
        print(f"\nGenerated files: {files}")


if __name__ == "__main__":
    main()
