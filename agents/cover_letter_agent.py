"""
Cover Letter Agent - Generates tailored cover letters.
Uses Claude Sonnet 4.5 for high-quality writing.
"""
import json
from pathlib import Path
from datetime import datetime
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Pt
from services.claude_client import ClaudeClient


class CoverLetterAgent:
    """Agent for generating tailored cover letters."""

    def __init__(self, profile_path: str = "data/candidate_profile.json"):
        self.profile_path = profile_path
        self.candidate_profile = self._load_profile()
        # COST OPTIMIZATION: Switch to Opus 4.5 (500K/min limit vs Sonnet's 30K/min)
        self.client = ClaudeClient(model="claude-opus-4-5")
        # IMPROVEMENT #3: Critic also uses Opus for consistency (same rate limit pool)
        self.critic_client = ClaudeClient(model="claude-opus-4-5")

    def _load_profile(self) -> Dict[str, Any]:
        """Load candidate profile."""
        with open(self.profile_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def generate_cover_letter(
        self,
        job: Dict[str, Any],
        match_report: Dict[str, Any],
        tailored_resume: Optional[str] = None,
        output_dir: str = "artifacts/cover_letters"
    ) -> Dict[str, str]:
        """Generate a tailored cover letter for a specific job."""
        print(f"\nGenerating cover letter for: {job['title']} at {job['company']}")
        print("-" * 80)

        cover_letter = self._generate_cover_letter_content(job, match_report, tailored_resume)

        job_hash = job.get('job_hash', 'unknown')
        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        md_file = output_path / f"cover_letter_{job_hash}.md"
        docx_file = output_path / f"cover_letter_{job_hash}.docx"

        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(cover_letter)

        self._create_docx(cover_letter, docx_file, job)

        print(f"  [OK] Markdown: {md_file}")
        print(f"  [OK] DOCX: {docx_file}")

        return {
            'markdown': str(md_file),
            'docx': str(docx_file)
        }

    def _generate_cover_letter_content(
        self,
        job: Dict[str, Any],
        match_report: Dict[str, Any],
        tailored_resume: Optional[str]
    ) -> str:
        """
        Generate cover letter content using Claude with quality checks.

        IMPROVEMENT #3: Implements Critic-Actor Loop
        - Actor (Sonnet): Generate initial draft
        - Critic (Haiku): Find 2 weakest sentences
        - Actor (Sonnet): Revise based on critique
        Total: max 3 API calls
        """
        # Extract company name - try multiple sources
        company_name = self._extract_company_name(job)

        prompt = self._build_cover_letter_prompt(job, match_report, tailored_resume, company_name)
        system_prompt = self._build_cover_letter_system_prompt()

        # ACTOR: Generate initial draft
        print(f"  [ACTOR] Generating initial draft...", flush=True)
        cover_letter_draft = self.client.generate(
            prompt=prompt,
            system=system_prompt,
            max_tokens=2048
        )

        # IMPROVEMENT #3: CRITIC-ACTOR LOOP
        # CRITIC: Analyze the draft for weaknesses
        print(f"  [CRITIC] Analyzing draft for weak sentences...", flush=True)
        critique = self._run_critic(cover_letter_draft, job['title'], company_name)

        job_hash = job.get('job_hash', 'unknown')

        # Save critique to file
        critique_file = Path(f"artifacts/cover_letters/critique_{job_hash}.json")
        critique_file.parent.mkdir(parents=True, exist_ok=True)
        with open(critique_file, 'w', encoding='utf-8') as f:
            json.dump(critique, f, indent=2)
        print(f"  [CRITIC] Saved critique to {critique_file}", flush=True)

        # ACTOR: Revise based on critique (if weaknesses found)
        if critique.get('weak_sentences') and len(critique['weak_sentences']) > 0:
            print(f"  [ACTOR] Revising based on critique ({len(critique['weak_sentences'])} weak sentences)...", flush=True)
            cover_letter = self._revise_with_critique(cover_letter_draft, critique, prompt, system_prompt)
        else:
            print(f"  [CRITIC] No significant weaknesses found - using original draft", flush=True)
            cover_letter = cover_letter_draft

        # Run self-check (existing quality check)
        cover_letter = self._run_quality_self_check(cover_letter, company_name, job['title'])

        return cover_letter.strip()

    def _extract_company_name(self, job: Dict[str, Any]) -> str:
        """Extract company name from job data or description."""
        company = job.get('company', 'Unknown Company')

        # If company is unknown/generic, try to extract from description
        if company in ['Unknown Company', 'None', 'Unknown', 'N/A', None]:
            raw_desc = job.get('raw_description', '')
            if raw_desc:
                # Ask Claude to extract company name
                try:
                    extract_prompt = f"""Extract the company name from this job description.
Return ONLY the company name, nothing else.
If no company name found, return "Unknown Company".

Job Description:
{raw_desc[:1000]}

Company name:"""
                    company = self.client.generate(
                        prompt=extract_prompt,
                        system="Extract only the company name from the text.",
                        max_tokens=50
                    ).strip()

                    if company and company not in ['Unknown Company', 'None', 'N/A', '']:
                        return company
                except:
                    pass

        return company if company else 'Unknown Company'

    def _run_critic(self, cover_letter: str, role: str, company: str) -> Dict[str, Any]:
        """
        IMPROVEMENT #3: Critic reviews the draft and identifies weak sentences.
        Uses Claude Haiku for cost-effectiveness.

        Returns:
            {
                "weak_sentences": ["sentence 1", "sentence 2"],
                "improvement_suggestions": ["suggestion 1", "suggestion 2"]
            }
        """
        critic_prompt = f"""You are a harsh critic reviewing this cover letter for a {role} role at {company}.

Your task: Find the 2 WEAKEST sentences - ones that are:
- Generic (could apply to any company/role)
- Vague or lack specifics
- Use clichés or buzzwords
- Don't add value or substance

Cover Letter:
{cover_letter}

Return JSON with exactly 2 weaknesses:
{{
    "weak_sentences": ["exact weak sentence 1", "exact weak sentence 2"],
    "improvement_suggestions": [
        "Specific improvement for sentence 1",
        "Specific improvement for sentence 2"
    ]
}}

If you find fewer than 2 weak sentences, return empty arrays."""

        try:
            critique = self.critic_client.generate_json(
                prompt=critic_prompt,
                system="You are a harsh but constructive critic of cover letters. Be specific and actionable in your feedback.",
                max_tokens=1024
            )

            # Log critique summary
            if critique.get('weak_sentences'):
                for idx, sentence in enumerate(critique['weak_sentences'][:2], 1):
                    print(f"    Weak sentence {idx}: {sentence[:60]}...", flush=True)

            return critique

        except Exception as e:
            print(f"  [WARNING] Critic failed: {e}", flush=True)
            return {'weak_sentences': [], 'improvement_suggestions': []}

    def _revise_with_critique(
        self,
        original_draft: str,
        critique: Dict[str, Any],
        original_prompt: str,
        system_prompt: str
    ) -> str:
        """
        IMPROVEMENT #3: Revise the cover letter based on critic's feedback.
        Only fixes the identified weak sentences.
        """
        weak_sentences = critique.get('weak_sentences', [])
        suggestions = critique.get('improvement_suggestions', [])

        # Build revision instructions
        revision_instructions = "Fix these specific weaknesses:\n"
        for idx, (weak, suggestion) in enumerate(zip(weak_sentences, suggestions), 1):
            revision_instructions += f"{idx}. Replace: \"{weak[:100]}...\"\n"
            revision_instructions += f"   With: {suggestion}\n\n"

        revision_prompt = f"""Revise this cover letter. Fix ONLY the weak sentences identified below.
Keep everything else IDENTICAL - do not change sentences that are already good.

{revision_instructions}

Original Cover Letter:
{original_draft}

CRITICAL: Output ONLY the revised cover letter text. Do NOT include your reasoning or explanation.

Revised cover letter:"""

        try:
            revised = self.client.generate(
                prompt=revision_prompt,
                system="You are a precise editor. Change only what was requested, keep everything else identical.",
                max_tokens=2048
            )

            return revised.strip()

        except Exception as e:
            print(f"  [WARNING] Revision failed: {e}", flush=True)
            return original_draft

    def _run_quality_self_check(self, cover_letter: str, company_name: str, role_title: str) -> str:
        """Run quality self-check and rewrite if needed."""
        banned_phrases = [
            "I am excited to",
            "I would love to",
            "passionate about",
            "I believe I would be a great fit",
            "I am writing to apply",
            "hit the ground running",
            "fast-paced environment"
        ]

        check_prompt = f"""Review this cover letter and check:
1. Does it contain any of these BANNED phrases: {', '.join(banned_phrases)}?
2. Does it mention the company name "{company_name}" specifically (if company is known)?
3. Does it open with a domain insight rather than "I am applying"?
4. Does it mention the role "{role_title}" specifically?

Cover Letter:
{cover_letter}

Return JSON:
{{
    "has_banned_phrases": true/false,
    "mentions_company": true/false,
    "has_domain_insight_opening": true/false,
    "mentions_role": true/false,
    "needs_rewrite": true/false,
    "reason": "explanation if rewrite needed"
}}"""

        try:
            check_result = self.client.generate_json(
                prompt=check_prompt,
                system="You are a strict cover letter reviewer.",
                max_tokens=500
            )

            # Skip company/role checks if company is unknown
            if company_name in ['Unknown Company', 'None', 'Unknown']:
                check_result['mentions_company'] = True

            if check_result.get('needs_rewrite') or check_result.get('has_banned_phrases'):
                print(f"  [WARNING] Cover letter failed self-check: {check_result.get('reason', 'Quality issues')}", flush=True)
                print(f"  [REWRITE] Regenerating cover letter...", flush=True)

                # Regenerate with stricter instructions
                rewrite_prompt = f"""The previous cover letter had issues: {check_result.get('reason')}

REWRITE COMPLETELY. Do not reuse any sentences from before.

Requirements:
- NO banned phrases: {', '.join(banned_phrases)}
- MUST open with domain/technical insight
- MUST mention company: {company_name} (if not Unknown Company)
- MUST mention role: {role_title}
- Use specific technical details from PersonaQuery and HealthEcho projects

Previous version (DO NOT COPY):
{cover_letter}

Write the NEW version now:"""

                cover_letter = self.client.generate(
                    prompt=rewrite_prompt,
                    system=self._build_cover_letter_system_prompt(),
                    max_tokens=2048
                )

        except Exception as e:
            print(f"  [WARNING] Self-check failed: {e}", flush=True)

        return cover_letter

    def _build_cover_letter_system_prompt(self) -> str:
        """Build system prompt for cover letter generation."""
        return """You are an expert cover letter writer specializing in thoughtful, engineer-focused letters.

STRUCTURE (exactly 4 paragraphs + sign-off):

**Paragraph 1 - Insight-first opening (2-3 sentences)**:
Start with a sharp observation about what makes this domain/problem genuinely hard or interesting. NOT "I am applying for...". End by connecting that insight to why this specific company's approach resonates with you.

**Paragraph 2 - Honest transparency (2-3 sentences)**:
If there's a skill gap between profile and JD, acknowledge it honestly but immediately reframe to what IS directly relevant. Be confident, not apologetic. This shows maturity and self-awareness.

**Paragraph 3 - Project specifics (4-5 sentences)**:
Describe PersonaQuery and HealthEcho with EXACT technical details from profile:
- PersonaQuery: evaluation harness, golden set for regression testing retrieval quality, citation guardrails, hallucination mitigation, p95 latency benchmarking, agent-assisted workflow, Groq (Llama-3.3-70B), LangChain, LlamaIndex
- HealthEcho: threshold tuning for ~40% false positive reduction, FastAPI inference layer, observability logging, anomaly detection pipeline, batch scoring

**Paragraph 4 - Company-specific closing (2-3 sentences)**:
Reference the exact problem the company is solving (from JD). Mention Sydney location and availability. Invite informal conversation.

**Sign-off**: "Thank you for your time and consideration.\n\nSincerely,\nRudra Joshi"

STRICTLY BANNED PHRASES (will fail if used):
❌ "I am excited to"
❌ "I would love to"
❌ "passionate about"
❌ "I believe I would be a great fit"
❌ "I am writing to apply"
❌ "hit the ground running"
❌ "fast-paced environment"

TONE RULES:
✓ Write like a thoughtful engineer discussing technical work
✓ Use specific technical vocabulary naturally (RAG, p95 latency, false positive rate)
✓ Direct, confident sentences
✓ Maximum 380 words
✓ ALWAYS mention company name in Para 4
✓ ALWAYS mention specific role title in Para 4
✓ ALWAYS mention Sydney in Para 4

OUTPUT FORMAT:
Just the cover letter text (4 paragraphs + sign-off), no additional commentary."""

    def _build_cover_letter_prompt(
        self,
        job: Dict[str, Any],
        match_report: Dict[str, Any],
        tailored_resume: Optional[str],
        company_name: str
    ) -> str:
        """Build prompt for cover letter generation."""
        profile_json = json.dumps(self.candidate_profile, indent=2)

        # Extract key fields
        role_title = job['title']
        raw_description = job.get('raw_description', '')[:2000]
        required_skills = job.get('required_skills', [])
        responsibilities = job.get('responsibilities', [])
        missing_skills = match_report.get('missing_skills', [])

        prompt = f"""BEFORE WRITING, ANSWER THESE:
(a) What specific hard problem does {company_name} solve? (infer from JD)
(b) Which 2 of Rudra's projects (PersonaQuery, HealthEcho) map most directly to the JD requirements?
(c) What skill gap (if any) should be acknowledged honestly?

JOB DETAILS:
Company: {company_name}
Role: {role_title}
Location: {job['location']}
Seniority: {job.get('seniority_level', 'N/A')}

Required Skills: {', '.join(required_skills[:10])}
Missing Skills: {', '.join(missing_skills[:5])}

Key Responsibilities:
{chr(10).join(f'- {r}' for r in responsibilities[:5])}

FULL JOB DESCRIPTION:
{raw_description}

MATCH ANALYSIS:
Score: {match_report['score']}/100
Top Matching Skills: {', '.join(match_report['matching_skills'][:5])}
Matching Reasons: {match_report['reasons'][:300]}

CANDIDATE PROFILE:
{profile_json}
"""

        if tailored_resume:
            prompt += f"\n\nTAILORED RESUME (for reference):\n{tailored_resume[:1000]}..."

        prompt += """

TASK:
First, think through these questions (DO NOT include your answers in the output):
(a) What specific hard problem does this company solve?
(b) Which 2 of Rudra's projects map to the JD?
(c) What skill gap should be acknowledged?

Then write ONLY the 4-paragraph cover letter (max 380 words):

Para 1: Start with insight about what makes this domain/problem hard, then connect to why this company
Para 2: Honest acknowledgment of any gap + immediate reframe to relevant strengths
Para 3: PersonaQuery and HealthEcho project specifics (evaluation harness, 40% FP reduction, p95 latency, FastAPI, etc.)
Para 4: Reference company's specific problem from JD + mention Sydney + invite conversation

End with: "Thank you for your time and consideration.\n\nSincerely,\nRudra Joshi"

CRITICAL: Output ONLY the cover letter text. Do NOT include your pre-writing analysis, thinking, or answers to (a)(b)(c).

Remember:
- NO "I am excited to" or "I would love to" or "passionate about"
- Use technical vocabulary naturally
- MUST mention company name and role in Para 4 (if known)
- MUST mention Sydney in Para 4"""

        return prompt

    def _create_docx(self, content: str, output_file: Path, job: Dict[str, Any]):
        """Convert cover letter to DOCX format."""
        doc = Document()

        # Set standard margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)

        # Add header with contact info
        header_para = doc.add_paragraph()
        header_para.add_run("Rudra Joshi\n").bold = True
        header_para.add_run(f"{self.candidate_profile.get('location', 'Sydney, NSW, Australia')}\n")

        if self.candidate_profile.get('education'):
            edu = self.candidate_profile['education'][0]
            header_para.add_run(f"{edu.get('institution', 'University of Wollongong')}\n")

        header_para.space_after = Pt(12)

        # Add date
        date_para = doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        date_para.space_after = Pt(12)

        # Add recipient
        recipient_para = doc.add_paragraph()
        recipient_para.add_run("Hiring Manager\n").bold = True
        recipient_para.add_run(f"{job['company']}\n")
        recipient_para.add_run(f"{job['location']}\n")
        recipient_para.space_after = Pt(12)

        # Add salutation
        salutation = doc.add_paragraph(f"Dear Hiring Manager,")
        salutation.space_after = Pt(12)

        # Add cover letter body
        paragraphs = content.split('\n\n')
        for para_text in paragraphs:
            if para_text.strip():
                para = doc.add_paragraph(para_text.strip())
                para.space_after = Pt(12)
                for run in para.runs:
                    run.font.size = Pt(11)

        # Add closing
        closing_para = doc.add_paragraph()
        closing_para.space_after = Pt(6)
        closing_para.add_run("Sincerely,\n\n")
        closing_para.add_run("Rudra Joshi").bold = True

        doc.save(output_file)


def main():
    """Example usage."""
    agent = CoverLetterAgent()

    # Load a shortlisted job
    with open('data/jobs_shortlisted.json', 'r') as f:
        shortlisted = json.load(f)

    if shortlisted:
        top_job = shortlisted[0]
        files = agent.generate_cover_letter(
            job=top_job['job'],
            match_report=top_job['match']
        )
        print(f"\nGenerated files: {files}")


if __name__ == "__main__":
    main()
